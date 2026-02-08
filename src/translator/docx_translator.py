import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.table import Table as DocxTable
from copy import deepcopy
from .base import BaseTranslator
import logging

logger = logging.getLogger(__name__)

WML_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class DocxTranslator(BaseTranslator):
    """Translator for DOCX files with formatting preservation."""

    def translate(self, input_path, output_path, target_language):
        """Translate a DOCX file while preserving formatting."""
        doc = Document(input_path)
        new_doc = Document()

        # Copy styles and formatting
        self._copy_document_properties(doc, new_doc)

        # Copy numbering definitions for bullets and numbered lists
        self._copy_numbering_definitions(doc, new_doc)

        logger.info(f"Starting DOCX translation to {target_language}")
        logger.info(f"Using model: {self.provider}/{self.model}")

        if self.progress_callback:
            self.progress_callback(5, "Preparing document...")

        # Process document body elements in order (paragraphs and tables interleaved)
        para_map = {para._element: para for para in doc.paragraphs}
        table_map = {table._element: table for table in doc.tables}

        body_elements = []
        for element in doc.element.body:
            if element.tag.endswith('p') and element in para_map:
                body_elements.append(('paragraph', para_map[element]))
            elif element.tag.endswith('tbl') and element in table_map:
                body_elements.append(('table', table_map[element]))

        total_elements = len(body_elements)
        BATCH_SIZE = 20

        # Batch paragraphs, but process tables immediately when encountered
        i = 0
        while i < total_elements:
            # Collect batch of consecutive paragraphs
            batch_paragraphs = []
            batch_elements = []  # Store (element_type, element, text, is_empty)

            batch_end = min(i + BATCH_SIZE, total_elements)
            for j in range(i, batch_end):
                elem_type, elem = body_elements[j]

                if elem_type == 'paragraph':
                    tagged_text, has_tags = self._runs_to_tagged_text(elem)
                    if tagged_text.strip():
                        batch_paragraphs.append(tagged_text)
                        batch_elements.append(('paragraph', elem, tagged_text, False, has_tags))
                    else:
                        batch_elements.append(('paragraph', elem, "", True, False))
                elif elem_type == 'table':
                    # Stop batching when we hit a table
                    batch_end = j
                    break

            # Translate and add batched paragraphs BEFORE checking for table
            if batch_elements:  # Changed from batch_paragraphs to batch_elements
                if batch_paragraphs:  # Only translate if there are non-empty paragraphs
                    logger.debug(f"Translating elements {i+1}-{batch_end}/{total_elements}...")

                    if self.progress_callback:
                        progress = 5 + int((i / total_elements) * 90)
                        self.progress_callback(progress, f"Translating elements {i+1}-{batch_end}/{total_elements}...")

                    translations = self.translate_batch(batch_paragraphs, target_language)
                    translation_iter = iter(translations)

                    # Apply translations in order
                    for elem_type, elem, text, is_empty, has_tags in batch_elements:
                        if is_empty:
                            new_doc.add_paragraph()
                        else:
                            translated_text = next(translation_iter)
                            new_paragraph = new_doc.add_paragraph()
                            if has_tags:
                                self._apply_paragraph_with_tags(elem, new_paragraph, translated_text)
                            else:
                                self._copy_paragraph_format(elem, new_paragraph, translated_text)
                else:
                    # Add empty paragraphs
                    for elem_type, elem, text, is_empty, has_tags in batch_elements:
                        new_doc.add_paragraph()

            # Move to next position
            i = batch_end

            # Now check if there's a table at the current position
            if i < total_elements and body_elements[i][0] == 'table':
                table = body_elements[i][1]
                logger.debug(f"Translating table at position {i+1}/{total_elements}...")

                if self.progress_callback:
                    progress = 5 + int((i / total_elements) * 90)
                    self.progress_callback(progress, f"Translating table at position {i+1}/{total_elements}...")

                self._translate_and_add_table(table, new_doc, target_language)
                i += 1

        if self.progress_callback:
            self.progress_callback(95, "Saving document...")

        new_doc.save(output_path)
        logger.info(f"Translation complete! Saved to {output_path}")

    def _copy_document_properties(self, source_doc, target_doc):
        """Copy document-level properties."""
        # Copy sections (page layout, margins, etc.)
        for section in source_doc.sections:
            target_section = target_doc.sections[-1] if target_doc.sections else target_doc.add_section()
            target_section.page_height = section.page_height
            target_section.page_width = section.page_width
            target_section.left_margin = section.left_margin
            target_section.right_margin = section.right_margin
            target_section.top_margin = section.top_margin
            target_section.bottom_margin = section.bottom_margin

    def _copy_numbering_definitions(self, source_doc, target_doc):
        """Copy numbering definitions from source to target document."""
        try:
            # Access the numbering part of the document
            source_numbering_part = source_doc.part.numbering_part
            if source_numbering_part is None:
                return  # No numbering in source document

            # Copy the entire numbering element to the target document
            if not hasattr(target_doc.part, 'numbering_part') or target_doc.part.numbering_part is None:
                # Add numbering part to target by copying from source
                target_doc.part._element.body.addnext(deepcopy(source_numbering_part._element))
            else:
                # Replace existing numbering definitions
                target_doc.part.numbering_part._element.clear()
                for child in source_numbering_part._element:
                    target_doc.part.numbering_part._element.append(deepcopy(child))
        except Exception as e:
            logger.warning(f"Could not copy numbering definitions: {e}")
            # Continue without numbering definitions - paragraphs will still work but may have incorrect formatting

    def _copy_paragraph_properties(self, source_para, target_para):
        """Copy paragraph-level formatting (style, alignment, spacing, numbering)."""
        if source_para.style:
            try:
                target_para.style = source_para.style.name
            except (KeyError, ValueError):
                pass

        target_para.alignment = source_para.alignment
        target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
        target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
        target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
        target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
        target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
        target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing

        source_pPr = source_para._element.pPr
        if source_pPr is not None:
            source_numPr = source_pPr.find(f'.//{WML_NS}numPr')
            if source_numPr is not None:
                target_pPr = target_para._element.get_or_add_pPr()
                existing_numPr = target_pPr.find(f'.//{WML_NS}numPr')
                if existing_numPr is not None:
                    target_pPr.remove(existing_numPr)
                target_pPr.append(deepcopy(source_numPr))

    def _strip_formatting_tags(self, text):
        """Remove any <b>, <i>, <u> tags (including malformed variants) that the LLM may have added."""
        # Well-formed tags: <b>, </b>, <i>, </i>, <u>, </u>
        text = re.sub(r'</?[biu]>', '', text)
        # Malformed tags from weaker models: <b=..., <i ..., </b=...>, etc.
        text = re.sub(r'</?[biu][^a-zA-Z>][^>]*>', '', text)
        # Unclosed malformed tags at end of text: <b=something (no closing >)
        text = re.sub(r'</?[biu][^a-zA-Z>][^>]*$', '', text)
        return text

    def _copy_font_properties(self, source_run, target_run):
        """Copy font properties (size, name, color) from source run to target run."""
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def _copy_paragraph_format(self, source_para, target_para, text):
        """Copy paragraph formatting and add text as a single uniformly-formatted run."""
        self._copy_paragraph_properties(source_para, target_para)
        text = self._strip_formatting_tags(text)

        if source_para.runs:
            first_run = source_para.runs[0]
            run = target_para.add_run(text)
            run.bold = first_run.bold
            run.italic = first_run.italic
            run.underline = first_run.underline
            self._copy_font_properties(first_run, run)
        else:
            target_para.add_run(text)

    def _runs_to_tagged_text(self, paragraph):
        """Convert paragraph runs to tagged text preserving inline formatting.

        Returns (text, has_tags). When all runs share the same formatting,
        returns plain text with has_tags=False.
        """
        runs = [r for r in paragraph.runs if r.text]
        if not runs:
            return paragraph.text, False

        # Merge consecutive runs with identical bold/italic/underline
        merged = []
        for run in runs:
            fmt = (bool(run.bold), bool(run.italic), bool(run.underline))
            if merged and merged[-1][0] == fmt:
                merged[-1] = (fmt, merged[-1][1] + run.text)
            else:
                merged.append((fmt, run.text))

        # If uniform formatting across all runs, no tags needed
        if len(set(m[0] for m in merged)) <= 1:
            return paragraph.text, False

        # Build tagged string
        tagged = ""
        for (bold, italic, underline), text in merged:
            segment = text
            if underline:
                segment = f"<u>{segment}</u>"
            if italic:
                segment = f"<i>{segment}</i>"
            if bold:
                segment = f"<b>{segment}</b>"
            tagged += segment

        return tagged, True

    def _parse_tagged_text(self, tagged_text):
        """Parse tagged text into list of (text, bold, italic, underline) tuples."""
        segments = []
        bold = italic = underline = False

        for part in re.split(r'(</?[biu]>)', tagged_text):
            if part == '<b>':
                bold = True
            elif part == '</b>':
                bold = False
            elif part == '<i>':
                italic = True
            elif part == '</i>':
                italic = False
            elif part == '<u>':
                underline = True
            elif part == '</u>':
                underline = False
            elif part:
                segments.append((part, bold, italic, underline))

        return segments

    def _apply_paragraph_with_tags(self, source_para, target_para, translated_text):
        """Apply translated tagged text to paragraph, creating runs per format segment."""
        self._copy_paragraph_properties(source_para, target_para)

        segments = self._parse_tagged_text(translated_text)

        # If parsing failed or leftover '<' in segments means malformed tags — fall back to plain text
        if not segments or any('<' in text for text, *_ in segments):
            cleaned = self._strip_formatting_tags(translated_text)
            self._copy_paragraph_format(source_para, target_para, cleaned)
            return

        base_run = source_para.runs[0] if source_para.runs else None

        for text, bold, italic, underline in segments:
            run = target_para.add_run(text)
            run.bold = True if bold else None
            run.italic = True if italic else None
            run.underline = True if underline else None
            if base_run:
                self._copy_font_properties(base_run, run)

    def _translate_and_add_table(self, table, new_doc, target_language):
        """Translate a table and add it to the new document.

        Deep-copies the source table element so that merges, borders, widths,
        row properties, and cell properties are all preserved.  Text is then
        replaced in-place in the copied structure.
        """
        # Deep copy preserves merges, borders, widths, row/cell properties
        new_tbl_element = deepcopy(table._element)
        # Insert before sectPr so the table stays in document order
        # (body.append would place it after sectPr, pushing it to the end)
        body = new_doc.element.body
        sectPr = body.find(f'{WML_NS}sectPr')
        if sectPr is not None:
            sectPr.addprevious(new_tbl_element)
        else:
            body.append(new_tbl_element)
        new_table = DocxTable(new_tbl_element, new_doc)

        # Collect unique non-empty cells for translation
        cell_texts = []
        cells_to_update = []  # (cell, has_tags, source_para)
        seen_cells = set()

        for row in new_table.rows:
            for cell in row.cells:
                cell_id = id(cell._element)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)

                if cell.text.strip():
                    # Use tagged text for cells with a single text paragraph
                    text_paras = [p for p in cell.paragraphs if p.text.strip()]
                    if len(text_paras) == 1:
                        tagged_text, has_tags = self._runs_to_tagged_text(text_paras[0])
                        cell_texts.append(tagged_text if has_tags else cell.text)
                        cells_to_update.append((cell, has_tags, text_paras[0]))
                    else:
                        cell_texts.append(cell.text)
                        cells_to_update.append((cell, False, None))

        # Translate in chunks to avoid overwhelming the LLM on large tables
        CELL_BATCH_SIZE = 15
        all_translated = []
        for start in range(0, len(cell_texts), CELL_BATCH_SIZE):
            chunk = cell_texts[start:start + CELL_BATCH_SIZE]
            all_translated.extend(self.translate_batch(chunk, target_language))

        # Replace text in-place, preserving run formatting
        for (cell, has_tags, source_para), translated_text in zip(cells_to_update, all_translated):
            if has_tags and source_para:
                # Save base run reference for font properties before modifying
                base_run = source_para.runs[0] if source_para.runs else None
                # Clear all runs across all paragraphs
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
                # Parse tags and create new runs in first paragraph
                segments = self._parse_tagged_text(translated_text)
                # Check for malformed tags: leftover '<' means the model corrupted the tags
                has_malformed = not segments or any('<' in text for text, *_ in segments)
                if not has_malformed and cell.paragraphs:
                    first_para = cell.paragraphs[0]
                    # Remove existing empty runs
                    for r_elem in first_para._element.findall(f'{WML_NS}r'):
                        first_para._element.remove(r_elem)
                    for text, bold, italic, underline in segments:
                        run = first_para.add_run(text)
                        run.bold = True if bold else None
                        run.italic = True if italic else None
                        run.underline = True if underline else None
                        if base_run:
                            self._copy_font_properties(base_run, run)
                else:
                    # Fallback: strip tags and put plain text
                    translated_text = self._strip_formatting_tags(translated_text)
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = translated_text
                    elif cell.paragraphs:
                        cell.paragraphs[0].add_run(translated_text)
            else:
                translated_text = self._strip_formatting_tags(translated_text)
                # Clear all existing run text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
                # Put translated text in first paragraph
                if cell.paragraphs:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = translated_text
                    else:
                        cell.paragraphs[0].add_run(translated_text)
